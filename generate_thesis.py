#!/usr/bin/env python3
"""
Diplomarbeit DOCX Generator
Liest k1-k7 Markdown-Dateien und generiert eine formatierte Diplomarbeit.
Verwendet Zitierregeln.docx als Style-Template.
"""
import re
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

BASE = os.path.dirname(os.path.abspath(__file__))

def read_md(filename):
    with open(os.path.join(BASE, filename), 'r', encoding='utf-8') as f:
        return f.read()

def add_page_break(doc):
    doc.add_page_break()

def set_paragraph_format(p, space_before=0, space_after=6, line_spacing=1.5):
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing

def add_title_page(doc):
    """Deckblatt / Titelseite"""
    for _ in range(6):
        doc.add_paragraph('')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('DIPLOMARBEIT')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    doc.add_paragraph('')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Digitales Zeiterfassungssystem')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Konzeption, Entwicklung und Cloud-Deployment\neiner webbasierten Arbeitszeiterfassung\nim MEVN-Stack')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    for _ in range(4):
        doc.add_paragraph('')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Ausgearbeitet von:')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Nawal Kayal\nAhmad Alalan')
    run.bold = True
    run.font.size = Pt(16)
    
    for _ in range(2):
        doc.add_paragraph('')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Schuljahr 2025/2026')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

def add_abstract(doc):
    """Kurzfassung / Abstract"""
    doc.add_heading('Kurzfassung', level=1)
    
    texts = [
        'Die vorliegende Diplomarbeit dokumentiert die Konzeption, Entwicklung und das Deployment eines vollständigen, cloud-basierten Zeiterfassungssystems. Das System wurde als moderne Single Page Application (SPA) im MEVN-Stack (MongoDB, Express.js, Vue.js 3, Node.js) realisiert und ermöglicht Unternehmen die rechtskonforme, digitale Erfassung von Arbeitszeiten ihrer Mitarbeitenden.',
        
        'Im Mittelpunkt der Arbeit steht ein Echtzeit-Stempelterminal, das über eine intuitive Weboberfläche bedient wird. Mitarbeitende können ihren Arbeitsbeginn, Pausen und das Arbeitsende mit einem einzigen Klick dokumentieren. Ein Live-Timer visualisiert die laufende Sitzungsdauer in Echtzeit. Ergänzend dazu bietet das System eine manuelle Nacherfassung für vergessene Buchungen, eine umfassende Urlaubs- und Abwesenheitsverwaltung mit mehrstufigem Genehmigungsworkflow sowie ein professionelles Abrechnungssystem (Billing) mit PDF-Export-Funktionalität.',
        
        'Die technische Architektur basiert auf einer strikten Client-Server-Trennung. Das Backend stellt eine RESTful API bereit, die durch JSON Web Tokens (JWT) und rollenbasierte Zugriffskontrolle (RBAC) abgesichert ist. Die Datenhaltung erfolgt in MongoDB Atlas, wobei Mongoose als Object Data Modeling Layer zum Einsatz kommt. Das Frontend nutzt Vue.js 3 mit der Composition API und wird über Vite als Build-Tool kompiliert. Für die Bereitstellung wird Render als Platform-as-a-Service (PaaS) verwendet.',
        
        'Besondere Herausforderungen bei der Implementierung umfassten die Echtzeit-Synchronisation des Live-Timers über Browser-Reloads hinweg, die performante Aggregation von Abrechnungsdaten mittels MongoDB Aggregation Pipelines sowie die Absicherung des SPA-Routings im Cloud-Deployment. Alle identifizierten Probleme wurden systematisch gelöst und sind in dieser Arbeit dokumentiert.',
        
        'Das fertige System ist unter einer festen URL im Internet erreichbar und wird produktiv genutzt. Es demonstriert die Praxistauglichkeit des gewählten Technologie-Stacks für mittelständische Unternehmensanwendungen.',
    ]
    for t in texts:
        p = doc.add_paragraph(t, style='Normal')
        set_paragraph_format(p, line_spacing=1.5)
    
    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Schlagwörter: ')
    run.bold = True
    run.font.size = Pt(11)
    p.add_run('Zeiterfassung, MEVN-Stack, Vue.js, Node.js, MongoDB, Cloud-Deployment, REST API, Single Page Application, JWT-Authentifizierung, Arbeitszeitgesetz').font.size = Pt(11)

def add_toc(doc):
    """Inhaltsverzeichnis (als Feld-Code)"""
    doc.add_heading('Inhaltsverzeichnis', level=1)
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    run2 = p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run2._r.append(instrText)
    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._r.append(fldChar2)
    run4 = p.add_run('[Inhaltsverzeichnis – Bitte mit Rechtsklick > „Felder aktualisieren" in Word generieren]')
    run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run4.font.size = Pt(10)
    run5 = p.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._r.append(fldChar3)

def add_list_of_figures(doc):
    """Abbildungsverzeichnis"""
    doc.add_heading('Abbildungsverzeichnis', level=1)
    figures = [
        ('Abbildung 1', 'Systemarchitektur des Zeiterfassungssystems (MEVN-Stack)'),
        ('Abbildung 2', 'Entity-Relationship-Diagramm der MongoDB Collections'),
        ('Abbildung 3', 'Dashboard-Ansicht mit Live-Stempelterminal'),
        ('Abbildung 4', 'Kalenderansicht mit FullCalendar-Integration'),
        ('Abbildung 5', 'Urlaubsantrag-Workflow (Statusdiagramm)'),
        ('Abbildung 6', 'Billing-Tabelle mit Soll-Ist-Vergleich'),
        ('Abbildung 7', 'Responsive Darstellung auf mobilen Endgeräten'),
        ('Abbildung 8', 'PDF-Export des Urlaubskontos'),
        ('Abbildung 9', 'Deployment-Architektur auf Render.com'),
        ('Abbildung 10', 'Sequenzdiagramm des Stempelvorgangs'),
    ]
    for label, desc in figures:
        p = doc.add_paragraph()
        run = p.add_run(f'{label}: ')
        run.bold = True
        p.add_run(desc)
        set_paragraph_format(p, space_after=3, line_spacing=1.15)

def add_abbreviations(doc):
    """Abkürzungsverzeichnis"""
    doc.add_heading('Abkürzungsverzeichnis', level=1)
    abbrevs = [
        ('API', 'Application Programming Interface'),
        ('CORS', 'Cross-Origin Resource Sharing'),
        ('CRUD', 'Create, Read, Update, Delete'),
        ('CSS', 'Cascading Style Sheets'),
        ('DOM', 'Document Object Model'),
        ('EuGH', 'Europäischer Gerichtshof'),
        ('ESM', 'ECMAScript Module'),
        ('HMR', 'Hot Module Replacement'),
        ('HR', 'Human Resources'),
        ('HTML', 'Hypertext Markup Language'),
        ('HTTP', 'Hypertext Transfer Protocol'),
        ('JSON', 'JavaScript Object Notation'),
        ('JWT', 'JSON Web Token'),
        ('KMU', 'Kleine und mittlere Unternehmen'),
        ('KPI', 'Key Performance Indicator'),
        ('MEVN', 'MongoDB, Express.js, Vue.js, Node.js'),
        ('NFC', 'Near Field Communication'),
        ('NoSQL', 'Not Only SQL'),
        ('ODM', 'Object Data Modeling'),
        ('PaaS', 'Platform as a Service'),
        ('RBAC', 'Role-Based Access Control'),
        ('REST', 'Representational State Transfer'),
        ('SaaS', 'Software as a Service'),
        ('SFC', 'Single File Component'),
        ('SPA', 'Single Page Application'),
        ('UI', 'User Interface'),
        ('URL', 'Uniform Resource Locator'),
        ('UTC', 'Coordinated Universal Time'),
        ('UX', 'User Experience'),
    ]
    table = doc.add_table(rows=1, cols=2, style='Grid Table 1 Light')
    hdr = table.rows[0].cells
    hdr[0].text = 'Abkürzung'
    hdr[1].text = 'Bedeutung'
    for cell in hdr:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for abbr, meaning in abbrevs:
        row = table.add_row().cells
        row[0].text = abbr
        row[1].text = meaning

def parse_markdown_to_doc(doc, md_text):
    """Parst Markdown-Text und fügt ihn dem Dokument hinzu."""
    lines = md_text.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    code_lang = ''
    
    while i < len(lines):
        line = lines[i]
        
        # Code-Block Start/Ende
        if line.strip().startswith('```'):
            if in_code_block:
                # Code-Block Ende
                code_text = '\n'.join(code_lines)
                if code_text.strip():
                    try:
                        p = doc.add_paragraph(style='Source Code Block')
                    except:
                        p = doc.add_paragraph()
                        for run in p.runs:
                            run.font.name = 'Consolas'
                            run.font.size = Pt(9)
                    p.clear()
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    pf = p.paragraph_format
                    pf.space_before = Pt(6)
                    pf.space_after = Pt(6)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lang = line.strip().replace('```', '')
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Leere Zeile
        if not line.strip():
            i += 1
            continue
        
        # Heading 1: # 
        if line.startswith('# '):
            title = line[2:].strip()
            # Entferne "/ Nawal /" etc.
            title = re.sub(r'\s*/\s*\w+(\s*&\s*\w+)?\s*/\s*', '', title).strip()
            doc.add_heading(title, level=1)
            i += 1
            continue
        
        # Heading 2: ##
        if line.startswith('## '):
            title = line[3:].strip()
            title = re.sub(r'\s*/\s*\w+(\s*&\s*\w+)?\s*/\s*', '', title).strip()
            doc.add_heading(title, level=2)
            i += 1
            continue
        
        # Heading 3: ###
        if line.startswith('### '):
            title = line[4:].strip()
            title = re.sub(r'\s*/\s*\w+(\s*&\s*\w+)?\s*/\s*', '', title).strip()
            doc.add_heading(title, level=3)
            i += 1
            continue
        
        # Horizontale Linie
        if line.strip() == '---':
            i += 1
            continue
        
        # Aufzählungsliste
        if line.strip().startswith('- **') or line.strip().startswith('- '):
            text = line.strip()[2:].strip()
            # Bold-Text verarbeiten
            p = doc.add_paragraph(style='List Paragraph')
            process_inline_formatting(p, text)
            set_paragraph_format(p, space_after=4, line_spacing=1.5)
            i += 1
            continue
        
        # Nummerierte Liste
        m = re.match(r'^(\d+)\.\s+\*\*(.+?)\*\*\s*(.*)', line.strip())
        if m:
            p = doc.add_paragraph(style='List Paragraph')
            run = p.add_run(f'{m.group(1)}. {m.group(2)}: ')
            run.bold = True
            if m.group(3):
                p.add_run(m.group(3))
            set_paragraph_format(p, space_after=4, line_spacing=1.5)
            i += 1
            continue
        
        m2 = re.match(r'^(\d+)\.\s+(.+)', line.strip())
        if m2:
            p = doc.add_paragraph(style='List Paragraph')
            process_inline_formatting(p, f'{m2.group(1)}. {m2.group(2)}')
            set_paragraph_format(p, space_after=4, line_spacing=1.5)
            i += 1
            continue
        
        # Normaler Absatz
        p = doc.add_paragraph(style='Normal')
        process_inline_formatting(p, line.strip())
        set_paragraph_format(p, space_after=6, line_spacing=1.5)
        i += 1

def process_inline_formatting(paragraph, text):
    """Verarbeitet **bold**, `code` und _italic_ Formatierung."""
    # Split nach Bold, Code, Italic patterns
    pattern = r'(\*\*.*?\*\*|`[^`]+`|_[^_]+_)'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x88, 0x22, 0x22)
        elif part.startswith('_') and part.endswith('_') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)

def add_bibliography(doc):
    """Literaturverzeichnis"""
    doc.add_heading('Literaturverzeichnis', level=1)
    
    sources = [
        '{1}  FLANAGAN, D., JavaScript: The Definitive Guide, 7. Auflage, Sebastopol, O\'Reilly Media, 2020, ISBN 978-1-491-95202-3',
        '{2}  YOU, E., Vue.js 3 – Official Documentation, Online: https://vuejs.org/guide/introduction.html, letzter Zugriff: März 2026',
        '{3}  MOZILLA DEVELOPER NETWORK (MDN), JavaScript Reference, Online: https://developer.mozilla.org/en-US/docs/Web/JavaScript, letzter Zugriff: März 2026',
        '{4}  MONGODB INC., MongoDB Manual – Aggregation Pipeline, Online: https://www.mongodb.com/docs/manual/core/aggregation-pipeline/, letzter Zugriff: März 2026',
        '{5}  EXPRESS.JS, Express – Node.js Web Application Framework, Online: https://expressjs.com/, letzter Zugriff: März 2026',
        '{6}  HAPI/JWT, JSON Web Tokens – Introduction, Online: https://jwt.io/introduction, letzter Zugriff: März 2026',
        '{7}  BROWN, E., Web Development with Node and Express, 2. Auflage, Sebastopol, O\'Reilly Media, 2019, ISBN 978-1-492-05351-4',
        '{8}  RENDER INC., Render Documentation – Deploy a Node Express App, Online: https://render.com/docs, letzter Zugriff: März 2026',
        '{9}  EUROPÄISCHER GERICHTSHOF (EuGH), Urteil C-55/18 (CCOO gegen Deutsche Bank), 14. Mai 2019',
        '{10} BOOTSTRAP, Bootstrap 5 Documentation, Online: https://getbootstrap.com/docs/5.3/, letzter Zugriff: März 2026',
        '{11} FULLCALENDAR, FullCalendar – JavaScript Calendar, Online: https://fullcalendar.io/docs, letzter Zugriff: März 2026',
        '{12} AXIOS, Axios – Promise based HTTP client, Online: https://axios-http.com/docs/intro, letzter Zugriff: März 2026',
        '{13} VITE, Vite – Next Generation Frontend Tooling, Online: https://vitejs.dev/guide/, letzter Zugriff: März 2026',
        '{14} MONGOOSE, Mongoose ODM v8.x Documentation, Online: https://mongoosejs.com/docs/guide.html, letzter Zugriff: März 2026',
        '{15} PROVOST, W., bcrypt – A Password Hashing Function, USENIX 1999',
        '{16} FIELDING, R., Architectural Styles and the Design of Network-based Software Architectures (REST), Dissertation, University of California, 2000',
        '{17} IETF, RFC 7519 – JSON Web Token (JWT), Online: https://www.rfc-editor.org/rfc/rfc7519, Mai 2015',
        '{18} HTML2PDF.JS, Client-side HTML-to-PDF rendering, Online: https://ekoopmans.github.io/html2pdf.js/, letzter Zugriff: März 2026',
    ]
    
    for s in sources:
        p = doc.add_paragraph(s, style='Normal')
        set_paragraph_format(p, space_after=8, line_spacing=1.3)

def add_eigenstaendigkeitserklaerung(doc):
    """Eigenständigkeitserklärung"""
    doc.add_heading('Eigenständigkeitserklärung', level=1)
    
    p = doc.add_paragraph('Wir erklären hiermit an Eides statt, dass wir die vorliegende Diplomarbeit selbstständig und ohne fremde Hilfe verfasst, andere als die angegebenen Quellen und Hilfsmittel nicht benutzt und die den benutzten Quellen wörtlich oder inhaltlich entnommenen Stellen als solche kenntlich gemacht haben.', style='Normal')
    set_paragraph_format(p, line_spacing=1.5)
    
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    p = doc.add_paragraph('Wien, am ___________________', style='Normal')
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    p = doc.add_paragraph('_______________________________          _______________________________', style='Normal')
    p = doc.add_paragraph('Nawal Kayal                                              Ahmad Alalan', style='Normal')

def main():
    print("=== Diplomarbeit DOCX Generator ===")
    print("Lade Template-Styles aus Zitierregeln.docx...")
    
    # Template laden (für Styles)
    doc = Document(os.path.join(BASE, 'Zitierregeln.docx'))
    
    # Bestehenden Inhalt löschen
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith('}sectPr'):
            continue
        body.remove(child)
    
    print("Template geladen. Generiere Diplomarbeit...")
    
    # 1. Deckblatt
    print("  -> Deckblatt...")
    add_title_page(doc)
    add_page_break(doc)
    
    # 2. Eigenständigkeitserklärung
    print("  -> Eigenständigkeitserklärung...")
    add_eigenstaendigkeitserklaerung(doc)
    add_page_break(doc)
    
    # 3. Kurzfassung
    print("  -> Kurzfassung / Abstract...")
    add_abstract(doc)
    add_page_break(doc)
    
    # 4. Inhaltsverzeichnis
    print("  -> Inhaltsverzeichnis...")
    add_toc(doc)
    add_page_break(doc)
    
    # 5. Abbildungsverzeichnis
    print("  -> Abbildungsverzeichnis...")
    add_list_of_figures(doc)
    add_page_break(doc)
    
    # 6. Abkürzungsverzeichnis
    print("  -> Abkürzungsverzeichnis...")
    add_abbreviations(doc)
    add_page_break(doc)
    
    # 7. Kapitel 1-7
    for i in range(1, 8):
        fname = f'diplomarbeit_k{i}.md'
        print(f"  -> Kapitel {i} ({fname})...")
        md = read_md(fname)
        parse_markdown_to_doc(doc, md)
        if i < 7:
            add_page_break(doc)
    
    # 8. Literaturverzeichnis
    print("  -> Literaturverzeichnis...")
    add_page_break(doc)
    add_bibliography(doc)
    
    # Speichern
    output = os.path.join(BASE, 'Diplomarbeit_Kayal_Alalan.docx')
    doc.save(output)
    print(f"\n✅ Diplomarbeit gespeichert: {output}")
    print(f"   Datei: {os.path.basename(output)}")

if __name__ == '__main__':
    main()
